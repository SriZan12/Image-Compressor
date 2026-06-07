from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/admin/Documents/Codex/2026-06-02/build-a-fully-functional-android-app")
OUTPUT = ROOT / "outputs" / "ImageCompressor-Technical-Documentation.docx"
SKILL_SCRIPTS = Path(
    "/Users/admin/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.601.10930/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF5D6"
WHITE = "FFFFFF"
BLACK = "000000"
CODE_FILL = "F6F8FA"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    props = element.get_or_add_pPr() if element.tag == qn("w:p") else element.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        data = kwargs.get(edge)
        if not data:
            continue
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        for key, value in data.items():
            node.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_keep_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_lines)
    p_pr.append(keep_next)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_bottom_border(paragraph, color=BLUE, size=8, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    source = doc.styles.add_style("Source Reference", 1)
    source.font.name = "Calibri"
    source.font.size = Pt(9)
    source.font.italic = True
    source.font.color.rgb = RGBColor.from_string(MUTED)
    source.paragraph_format.space_before = Pt(1)
    source.paragraph_format.space_after = Pt(6)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("IMAGE COMPRESSOR")
    set_run_font(left, size=9, color=DARK_BLUE, bold=True)
    right = header.add_run("   |   Technical Developer Handbook")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(footer)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_para(doc, text="", *, bold_prefix=None, italic=False, color=BLACK, after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=color)
        remaining = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(remaining, italic=italic, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic, color=color)
    return paragraph


def add_rich_para(doc, chunks, *, after=6, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(after)
    for text, options in chunks:
        run = paragraph.add_run(text)
        set_run_font(
            run,
            name=options.get("font", "Calibri"),
            size=options.get("size", 11),
            color=options.get("color", BLACK),
            bold=options.get("bold"),
            italic=options.get("italic"),
        )
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def new_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "7")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def add_number(doc, text, num_id):
    paragraph = doc.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(number)
    p_pr.insert(0, num_pr)
    return paragraph


def add_numbered_steps(doc, steps):
    num_id = new_decimal_numbering(doc)
    for step in steps:
        add_number(doc, step, num_id)


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        paragraph = doc.add_paragraph(style="Code Block")
        paragraph.paragraph_format.space_after = Pt(0)
        shade(paragraph._p, CODE_FILL)
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, name="Courier New", size=8.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_callout(doc, label, text, *, fill=CALLOUT, label_color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    shade(paragraph._p, fill)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, bold=True, color=label_color)
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)
    return paragraph


def add_source(doc, *references):
    paragraph = doc.add_paragraph(style="Source Reference")
    run = paragraph.add_run("Source reference: " + "; ".join(references))
    set_run_font(run, size=9, color=MUTED, italic=True)
    return paragraph


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].cells
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=font_size, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size, color=BLACK)
    apply_table_geometry(table, widths_dxa, table_width_dxa=sum(widths_dxa), indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_logic_section(doc, number, title, purpose, algorithm_steps, example, sources, notes=None):
    add_heading(doc, f"{number}. {title}", 2)
    add_rich_para(
        doc,
        [
            ("Purpose. ", {"bold": True, "color": DARK_BLUE}),
            (purpose, {}),
        ],
        after=5,
    )
    add_rich_para(doc, [("Algorithm. ", {"bold": True, "color": DARK_BLUE})], after=2)
    add_numbered_steps(doc, algorithm_steps)
    add_callout(doc, "Worked example", example, fill="EEF5FB")
    if notes:
        add_callout(doc, "Engineering note", notes, fill=CAUTION, label_color="7A5A00")
    add_source(doc, *sources)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("DEVELOPER HANDBOOK")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Image Compressor")
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Technical Documentation")
    set_run_font(run, size=18, color=DARK_BLUE, bold=True)

    subsubtitle = doc.add_paragraph()
    subsubtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subsubtitle.paragraph_format.space_after = Pt(20)
    run = subsubtitle.add_run("Architecture, logic, algorithms, and worked examples")
    set_run_font(run, size=12, color=MUTED, italic=True)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(22)
    add_bottom_border(rule, color=BLUE, size=8)

    metadata = [
        ("Project", "Image Compressor Android app"),
        ("Technology", "Kotlin, Jetpack Compose, Room, Coroutines, Flow, Coil"),
        ("Platform", "Android 8.0+ (minSdk 26), targetSdk 37"),
        ("Document scope", "Implemented source code in this workspace"),
        ("Document date", "June 2, 2026"),
    ]
    add_table(doc, ["Field", "Value"], metadata, [2100, 7260], font_size=10)
    add_callout(
        doc,
        "Privacy model",
        "All compression work is performed locally on the device. The app does not upload selected images to a server.",
        fill="EAF5EE",
        label_color="1F5B36",
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "How to Read This Handbook", 1)
    add_para(
        doc,
        "This handbook explains the implemented application as a working system. Each logic section includes "
        "the goal, execution steps, a concrete example, important caveats, and direct source-file references. "
        "The examples are intentionally numeric where possible so that the behavior can be checked by hand.",
    )
    add_callout(
        doc,
        "Notation",
        "Source references use repository-relative paths followed by line numbers, for example "
        "util/ImageCompressor.kt:145. Byte examples use 1 KB = 1024 bytes and 1 MB = 1024 KB.",
    )

    add_heading(doc, "Contents", 1)
    contents = [
        ("1", "Project overview and requirements mapping"),
        ("2", "Architecture and end-to-end runtime flow"),
        ("3", "Core data models and derived values"),
        ("4", "Image selection and metadata inspection logic"),
        ("5", "Resize, sampling, and memory-protection algorithms"),
        ("6", "Compression, target-size, and output-file algorithms"),
        ("7", "Batch orchestration, progress, and history persistence"),
        ("8", "Gallery saving, sharing, and Android-version handling"),
        ("9", "Compose UI, navigation, theme, and status handling"),
        ("10", "Privacy, permissions, failure handling, and known constraints"),
        ("11", "Build configuration, testing, and verification"),
    ("Appendix A", "Pseudocode reference"),
    ("Appendix B", "Source-file ownership map"),
    ("Appendix C", "Practical extension playbook"),
]
    add_table(doc, ["Section", "Coverage"], contents, [1700, 7660], font_size=10)

    add_heading(doc, "1. Project Overview", 1)
    add_para(
        doc,
        "Image Compressor is a privacy-first Android application for selecting one or more images, previewing "
        "their metadata, applying local compression settings, comparing the results, and saving or sharing the "
        "compressed files. The app is deliberately serverless: input images stay on the device.",
    )
    add_heading(doc, "1.1 Functional Coverage", 2)
    feature_rows = [
        ("Photo selection", "Single and multi-select through Android Photo Picker", "Implemented"),
        ("Preview", "Grid cards with thumbnail, file size, resolution, and format", "Implemented"),
        ("Compression controls", "Quality, target size, resize mode, aspect ratio, JPEG/PNG/WEBP", "Implemented"),
        ("Local processing", "Coroutine-based bitmap work on Dispatchers.IO", "Implemented"),
        ("Results", "Original/compressed previews, sizes, reduction percentage", "Implemented"),
        ("Export", "Gallery save and Android share sheet", "Implemented"),
        ("History", "Room records with timestamp, sizes, format, dimensions, and saved URI", "Implemented"),
        ("Privacy onboarding", "First-run local-processing explanation", "Implemented"),
        ("Themes", "System, light, dark, and Android 12+ dynamic color", "Implemented"),
    ]
    add_table(doc, ["Area", "Behavior", "Status"], feature_rows, [1900, 5860, 1600], font_size=9.2)
    add_source(doc, "README.md", "ui/ImageCompressorApp.kt", "util/ImageCompressor.kt")

    add_heading(doc, "2. Architecture and Runtime Flow", 1)
    add_para(
        doc,
        "The implementation follows MVVM with a Repository pattern. Compose renders immutable UI state. "
        "The ViewModel owns transitions and launches coroutines. The repository coordinates image utilities "
        "and Room. Utilities isolate Android bitmap, MediaStore, and share-intent details.",
    )
    add_heading(doc, "2.1 Layer Responsibilities", 2)
    layer_rows = [
        ("Compose UI", "Renders screens and dispatches user actions", "ui/ImageCompressorApp.kt"),
        ("ViewModel", "Owns screen state, progress, messages, and orchestration", "ui/ImageCompressorViewModel.kt"),
        ("Repository", "Coordinates inspection, compression, save, and Room history", "data/ImageRepository.kt"),
        ("Compression utility", "Decodes, resizes, rotates, encodes, and caches images", "util/ImageCompressor.kt"),
        ("Storage utilities", "Writes MediaStore files and launches share intents", "util/GallerySaver.kt; util/ShareUtils.kt"),
        ("Local persistence", "Room database and SharedPreferences-backed flows", "data/local/*; data/preferences/*"),
        ("Dependency container", "Builds and shares application-scoped dependencies", "ImageCompressorApplication.kt"),
    ]
    add_table(doc, ["Layer", "Responsibility", "Primary source"], layer_rows, [1700, 4260, 3400], font_size=9)

    add_heading(doc, "2.2 End-to-End Flow", 2)
    add_numbered_steps(doc, [
        "The user completes onboarding and opens the Home screen.",
        "The Photo Picker returns one or more content URIs.",
        "The repository attempts to retain URI read access and inspects metadata off the main thread.",
        "The user chooses quality, target size, resize mode, aspect-ratio behavior, and output format.",
        "The repository compresses each image sequentially on Dispatchers.IO and inserts a Room history row.",
        "The ViewModel updates progress after each completed image.",
        "The Results screen compares before and after values and offers save or share actions.",
        "Saving writes to MediaStore on Android 10+ or the public Pictures directory on Android 8/9.",
    ])
    add_callout(
        doc,
        "Example flow",
        "A user selects three photos, chooses 80% JPEG quality and a 500 KB target, then taps Compress Images. "
        "The progress state moves through 1/3, 2/3, and 3/3. Each output is cached, recorded in Room, shown in "
        "the Results screen, and saved only if the user taps Save to Gallery.",
        fill="EEF5FB",
    )
    add_source(doc, "ui/ImageCompressorApp.kt:114-232", "ui/ImageCompressorViewModel.kt:101-215", "data/ImageRepository.kt:43-81")

    add_heading(doc, "3. Core Data Models and Derived Values", 1)
    add_heading(doc, "3.1 Domain Models", 2)
    model_rows = [
        ("CompressionSettings", "User-selected quality, target bytes, resize mode, dimensions, aspect-ratio toggle, format"),
        ("SelectedImage", "Input URI plus inspected display name, bytes, resolution, and format"),
        ("CompressedImage", "Output URI/path, output bytes, dimensions, format, history row ID, optional saved URI"),
        ("CompressionHistoryEntity", "Room record for a completed compression result"),
        ("ImageCompressorUiState", "Single immutable snapshot consumed by the Compose UI"),
    ]
    add_table(doc, ["Model", "Purpose"], model_rows, [2400, 6960], font_size=9.5)

    add_logic_section(
        doc,
        "3.2",
        "Reduction Percentage",
        "Convert the before/after byte counts into a beginner-friendly percentage shown in Results and History.",
        [
            "If the original byte count is zero or unknown, return 0 to avoid division by zero.",
            "Compute 1 - (compressedBytes / originalBytes).",
            "Multiply by 100, round to the nearest integer, and clamp to the range 0..100.",
            "Clamping prevents a larger output file from displaying a negative reduction.",
        ],
        "For a 4 MB original and a 1 MB output: (1 - 1/4) * 100 = 75%. "
        "For a 100 KB original and a 150 KB output, the raw value is -50%, but the displayed value is clamped to 0%.",
        ["data/model/CompressionModels.kt:61-66", "ui/ImageCompressorApp.kt:695-703"],
    )

    add_logic_section(
        doc,
        "3.3",
        "Readable File Sizes",
        "Format raw byte counts into readable KB or MB labels for the UI.",
        [
            "Return 'Unknown size' for non-positive values.",
            "Divide bytes by 1024 to obtain KB.",
            "If KB is less than 1024, show one decimal place.",
            "Otherwise divide KB by 1024 and show MB with two decimal places.",
        ],
        "1024 bytes becomes 1.0 KB. 1,048,576 bytes becomes 1.00 MB. A provider that reports 0 bytes "
        "is displayed as Unknown size.",
        ["data/model/CompressionModels.kt:68-73", "data/model/CompressionModelsTest.kt:14-18"],
    )

    add_heading(doc, "4. Image Selection and Metadata Inspection", 1)
    add_logic_section(
        doc,
        "4.1",
        "Photo Picker Launch",
        "Select only the images the user explicitly chooses, without requesting broad media-library read access.",
        [
            "Register PickMultipleVisualMedia with a maximum of 30 images and PickVisualMedia for a single image.",
            "Launch each picker with ImageOnly so videos are excluded.",
            "Pass returned URIs to the ViewModel.",
            "Use AndroidX Photo Picker behavior so older supported devices can fall back to the system document picker.",
        ],
        "On Home, Select Images launches the multi-picker. In Preview, Add One launches the single-picker and "
        "Add Multiple launches the multi-picker. Selecting two new photos sends a two-URI list to addSelectedImages().",
        ["ui/ImageCompressorApp.kt:114-146", "AndroidManifest.xml:29-41"],
    )

    add_logic_section(
        doc,
        "4.2",
        "URI Retention and De-duplication",
        "Retain read access where the URI provider supports it and avoid duplicate cards when the same image is selected twice.",
        [
            "Call distinct() on the incoming URI list before inspecting it.",
            "Attempt takePersistableUriPermission() with read access inside runCatching.",
            "Continue even if a provider does not support persistable grants.",
            "After inspection, merge old and new selections and call distinctBy { id }, where id is uri.toString().",
        ],
        "If the existing selection contains content://photos/42 and the next picker response contains "
        "content://photos/42 plus content://photos/99, only two cards remain: 42 and 99.",
        ["data/ImageRepository.kt:43-54", "data/model/CompressionModels.kt:35-44", "ui/ImageCompressorViewModel.kt:101-126"],
        "Persistable URI grants are best-effort because Photo Picker providers and document providers do not all expose identical grant behavior.",
    )

    add_logic_section(
        doc,
        "4.3",
        "Metadata Query",
        "Show useful metadata before compression without decoding a full-resolution bitmap into memory.",
        [
            "Query OpenableColumns.DISPLAY_NAME and OpenableColumns.SIZE through ContentResolver.",
            "If the reported size is missing or non-positive, ask openAssetFileDescriptor() for the asset length.",
            "Decode bounds only to obtain pixel width and height.",
            "Read EXIF orientation and swap displayed width/height for 90-degree or 270-degree rotation.",
            "Resolve the MIME type and show its subtype in uppercase.",
        ],
        "A provider reports vacation.jpg as 3,145,728 bytes with encoded bounds 4000 x 3000 and EXIF rotation 90. "
        "The preview card shows 3.00 MB, 3000 x 4000, and JPEG.",
        ["util/ImageCompressor.kt:35-51", "util/ImageCompressor.kt:114-143", "util/ImageCompressor.kt:169-185"],
    )

    add_logic_section(
        doc,
        "4.4",
        "Bounds-Only Decode",
        "Read image dimensions safely before allocating pixel memory.",
        [
            "Create BitmapFactory.Options with inJustDecodeBounds = true.",
            "Open the content URI stream and fail clearly if the provider cannot open it.",
            "Call BitmapFactory.decodeStream(). In bounds-only mode, its return value is intentionally ignored.",
            "Validate options.outWidth and options.outHeight after decoding.",
        ],
        "For an 8000 x 6000 photo, decodeStream() returns no bitmap because inJustDecodeBounds is true, but "
        "options.outWidth becomes 8000 and options.outHeight becomes 6000. The app uses those options values.",
        ["util/ImageCompressor.kt:134-143"],
        "This distinction matters: treating the expected null bitmap return as an error would reject valid images. "
        "The implementation explicitly checks stream availability separately from populated bounds.",
    )

    add_heading(doc, "5. Resize, Sampling, and Memory Protection", 1)
    add_logic_section(
        doc,
        "5.1",
        "Percentage Resize",
        "Scale both dimensions by the user's chosen percentage while preserving the image proportions.",
        [
            "Clamp the percentage to 10..100.",
            "Convert it into a decimal scale by dividing by 100.",
            "Multiply width and height by that scale.",
            "Round each value and ensure each dimension is at least 1 pixel.",
        ],
        "A 4000 x 3000 image resized to 50% becomes 2000 x 1500. A 10% resize becomes 400 x 300.",
        ["util/ImageCompressor.kt:202-208", "ui/ImageCompressorViewModel.kt:150-152"],
    )

    add_logic_section(
        doc,
        "5.2",
        "Custom Resize With Aspect Ratio",
        "Fit the original image inside a custom width/height box without stretching it.",
        [
            "Parse the width and height text fields. If a value is missing, fall back to the source dimension.",
            "When aspect-ratio preservation is enabled, compute widthScale = requestedWidth / sourceWidth.",
            "Compute heightScale = requestedHeight / sourceHeight.",
            "Use the smaller scale so both output dimensions fit inside the requested box.",
            "When preservation is disabled, use the requested dimensions directly.",
        ],
        "A 4000 x 3000 image placed into a 1000 x 1000 box has widthScale = 0.25 and heightScale = 0.333. "
        "The smaller scale is 0.25, so the result is 1000 x 750. With preservation disabled, it becomes 1000 x 1000.",
        ["util/ImageCompressor.kt:209-219", "ui/ImageCompressorApp.kt:518-544"],
    )

    add_logic_section(
        doc,
        "5.3",
        "Pixel Budget Constraint",
        "Prevent unusually large requested dimensions from creating oversized in-memory bitmaps.",
        [
            "Calculate pixels = width * height.",
            "If pixels is at most 16,000,000, keep the dimensions.",
            "Otherwise compute scale = sqrt(16,000,000 / pixels).",
            "Multiply both dimensions by that scale to preserve proportions while approaching the pixel limit.",
        ],
        "An 8000 x 6000 request contains 48,000,000 pixels. scale = sqrt(16,000,000 / 48,000,000) = 0.577. "
        "The constrained dimensions are approximately 4619 x 3464, close to 16 million pixels.",
        ["util/ImageCompressor.kt:220-230", "util/ImageCompressor.kt:270-272"],
    )

    add_logic_section(
        doc,
        "5.4",
        "Sampled Bitmap Decode",
        "Reduce memory usage by decoding an appropriately sampled bitmap instead of always loading the full image.",
        [
            "Read source bounds first.",
            "Start inSampleSize at 1.",
            "Double the sample size while the next downsampled width and height are still at least the requested dimensions.",
            "Continue doubling if the decoded pixel count would remain above the 16-million-pixel memory budget.",
            "Decode the URI stream using the chosen inSampleSize and ARGB_8888.",
        ],
        "For an 8000 x 6000 image requested at 2000 x 1500, sample size progresses 1 -> 2 -> 4. "
        "A further jump to 8 would decode only 1000 x 750, below the request, so the selected sample is 4.",
        ["util/ImageCompressor.kt:145-167"],
        "Sampling is power-of-two based because BitmapFactory.inSampleSize is designed around efficient decoder sampling.",
    )

    add_logic_section(
        doc,
        "5.5",
        "EXIF-Aware Rotation",
        "Keep compressed output visually upright when the camera stored orientation as metadata rather than rotated pixels.",
        [
            "Read EXIF TAG_ORIENTATION from the selected URI stream.",
            "Map ROTATE_90, ROTATE_180, and ROTATE_270 to degree values.",
            "Create a Matrix rotation during bitmap decode.",
            "Recycle the pre-rotation bitmap if a new rotated bitmap is created.",
            "Write ORIENTATION_NORMAL to the exported file because its pixels are already upright.",
        ],
        "A portrait photo may be stored as 4032 x 3024 pixels plus ROTATE_90 metadata. The compressor rotates the pixels "
        "and exports an upright 3024 x 4032 image marked ORIENTATION_NORMAL.",
        ["util/ImageCompressor.kt:90-96", "util/ImageCompressor.kt:169-193"],
        "The implementation handles rotational EXIF values. Mirrored EXIF orientations are not transformed.",
    )

    add_heading(doc, "6. Compression and Output Algorithms", 1)
    add_logic_section(
        doc,
        "6.1",
        "Main Compression Pipeline",
        "Convert a SelectedImage and CompressionSettings object into a cached CompressedImage result.",
        [
            "Calculate requested output dimensions.",
            "Decode a sampled bitmap and apply EXIF rotation.",
            "Scale the decoded bitmap to the final requested dimensions.",
            "Flatten transparency onto white when exporting JPEG.",
            "Encode to bytes at the chosen quality.",
            "Apply target-size iteration when enabled.",
            "Write a uniquely named cache file, normalize EXIF orientation, recycle the bitmap, and return output metadata.",
        ],
        "A 4000 x 3000 PNG is configured for 50% resize and JPEG output at 80 quality. It is decoded efficiently, "
        "scaled to 2000 x 1500, flattened onto white if it has alpha, JPEG encoded, and written into cache/compressed/.",
        ["util/ImageCompressor.kt:53-112"],
    )

    add_logic_section(
        doc,
        "6.2",
        "JPEG Transparency Flattening",
        "Avoid unexpected dark or transparent backgrounds when an alpha-enabled input is converted to JPEG.",
        [
            "Check whether the requested output format is JPEG and whether the bitmap reports alpha.",
            "Create an ARGB_8888 destination bitmap with matching dimensions.",
            "Fill its canvas with white.",
            "Draw the original bitmap over the white background.",
            "Recycle the old bitmap and encode the flattened result.",
        ],
        "A transparent PNG logo converted to JPEG gets a white background. Without flattening, transparent pixels "
        "could be rendered unpredictably by the encoder or downstream viewer.",
        ["util/ImageCompressor.kt:239-248"],
    )

    add_logic_section(
        doc,
        "6.3",
        "Format-Specific Encoding",
        "Map the user's output choice to Android bitmap encoders while accounting for platform differences.",
        [
            "Use Bitmap.CompressFormat.JPEG for JPEG output.",
            "Use Bitmap.CompressFormat.PNG for PNG output.",
            "Use WEBP_LOSSY on Android 11+ and the older WEBP constant on earlier supported versions.",
            "Write encoded bytes into a ByteArrayOutputStream and fail if bitmap.compress() returns false.",
        ],
        "On Android 13, WEBP uses WEBP_LOSSY. On Android 9, the app falls back to the legacy WEBP encoder constant. "
        "PNG remains lossless, so the quality slider has less influence on PNG file size.",
        ["util/ImageCompressor.kt:250-263", "data/model/CompressionModels.kt:6-10"],
    )

    add_logic_section(
        doc,
        "6.4",
        "Target File Size Iteration",
        "Try to bring the output close to an optional target size without infinite looping.",
        [
            "Encode once using the configured quality.",
            "While encoded bytes exceed the target and attempts are below 40, continue adjusting.",
            "For JPEG or WEBP with quality above 10, reduce quality by 5 and encode again.",
            "For PNG, or after lossy quality reaches 10, reduce bitmap width and height to 90% of the prior values.",
            "Stop when bytes are within target, dimensions no longer change, or 40 attempts have run.",
        ],
        "A JPEG initially encodes to 820 KB at quality 80 with a 500 KB target. The loop retries at qualities "
        "75, 70, 65, and so on. If quality 10 is still too large, dimensions shrink to 90% and encoding continues.",
        ["util/ImageCompressor.kt:63-81", "util/ImageCompressor.kt:265-271"],
        "The target is best-effort. Some images cannot reach the requested size without excessive dimension loss, "
        "and the 40-attempt cap prevents unbounded work.",
    )

    add_logic_section(
        doc,
        "6.5",
        "Unique Cache Output and FileProvider URI",
        "Create shareable temporary result files without exposing raw filesystem paths to other apps.",
        [
            "Ensure cacheDir/compressed exists.",
            "Build a filename from a fixed prefix, current epoch milliseconds, an eight-character UUID suffix, and format extension.",
            "Write the encoded bytes to that file.",
            "Create a content URI using FileProvider and the app-specific authority.",
            "Return both the private file path for internal copying and the content URI for previews or sharing.",
        ],
        "One JPEG output might be named ImageCompressor_1780412812345_a1b2c3d4.jpg. Another compression in the same "
        "session receives a different timestamp or UUID suffix, avoiding collisions.",
        ["util/ImageCompressor.kt:83-112", "AndroidManifest.xml:19-27", "res/xml/file_paths.xml:1-6"],
    )

    add_heading(doc, "7. Batch Processing, Progress, and Persistence", 1)
    add_logic_section(
        doc,
        "7.1",
        "Sequential Batch Compression",
        "Compress multiple images off the main thread while producing predictable progress updates and Room records.",
        [
            "Move repository work into withContext(Dispatchers.IO).",
            "Iterate through selected images with mapIndexed.",
            "Compress one image.",
            "Insert its completed result into Room and capture the generated history ID.",
            "Invoke onProgress(index + 1, total).",
            "Return the CompressedImage with its history ID attached.",
        ],
        "For three selected photos, the repository emits progress after each output: completed=1,total=3; "
        "completed=2,total=3; completed=3,total=3. The UI fractions are 0.333, 0.667, and 1.0.",
        ["data/ImageRepository.kt:56-81", "ui/ImageCompressorViewModel.kt:174-215"],
        "Compression is sequential rather than parallel. This is conservative for memory use because several large bitmaps are not decoded at the same time.",
    )

    add_logic_section(
        doc,
        "7.2",
        "Progress Fraction",
        "Convert completed and total counts into a value suitable for Material 3 LinearProgressIndicator.",
        [
            "If total is zero, return 0f.",
            "Otherwise divide completed.toFloat() by total.",
            "Pass the fraction to the Results loading card while isCompressing is true.",
        ],
        "During the second image of a four-image batch, completed=2 and total=4. fraction = 2 / 4 = 0.5, "
        "so the progress bar displays 50%.",
        ["ui/ImageCompressorViewModel.kt:32-35", "ui/ImageCompressorApp.kt:577-586"],
    )

    add_logic_section(
        doc,
        "7.3",
        "Room History Flow",
        "Persist completed compression summaries locally and keep the History UI reactive.",
        [
            "Store each completion as CompressionHistoryEntity in the compression_history table.",
            "Observe rows ordered by createdAt descending as Flow<List<CompressionHistoryEntity>>.",
            "Collect the Flow in the ViewModel and copy the latest list into UI state.",
            "Update savedPath after a successful gallery save.",
            "Support deleting one row or clearing all rows.",
        ],
        "After compressing beach.jpg, Room stores originalSizeBytes=3145728 and compressedSizeBytes=524288. "
        "When the user saves it, updateSavedPath() adds the returned MediaStore content URI. History refreshes automatically.",
        ["data/local/CompressionHistoryEntity.kt:6-19", "data/local/CompressionHistoryDao.kt:9-23", "data/ImageRepository.kt:83-95"],
    )

    add_heading(doc, "8. Gallery Save, Sharing, and Android Versions", 1)
    add_logic_section(
        doc,
        "8.1",
        "Scoped MediaStore Save on Android 10+",
        "Publish a compressed file to the gallery without requesting legacy broad storage permission.",
        [
            "Insert an Images.Media row with display name, MIME type, Pictures/Image Compressor relative path, and IS_PENDING=1.",
            "Open the destination output stream and copy the private cache file.",
            "Clear the pending marker by updating IS_PENDING=0.",
            "If copying fails, delete the incomplete MediaStore row and rethrow the error.",
        ],
        "On Android 14, saving ImageCompressor_...jpg inserts a row under Pictures/Image Compressor, copies bytes, "
        "then makes the row visible by clearing IS_PENDING. No runtime storage prompt is shown.",
        ["util/GallerySaver.kt:22-56"],
    )

    add_logic_section(
        doc,
        "8.2",
        "Legacy Save on Android 8 and 9",
        "Support API 26-28 devices that predate scoped storage.",
        [
            "Check WRITE_EXTERNAL_STORAGE permission before writing.",
            "Create the public Pictures/Image Compressor directory.",
            "Copy the cache file into the public directory.",
            "Ask MediaScannerConnection to index the file.",
            "Return the indexed URI, or a file URI fallback if scanning returns null.",
        ],
        "On Android 9, the first save triggers the runtime storage prompt. After approval, the cache JPEG is copied "
        "to Pictures/Image Compressor and scanned so gallery apps can display it.",
        ["ui/ImageCompressorApp.kt:122-140", "util/GallerySaver.kt:58-82", "AndroidManifest.xml:5-7"],
    )

    add_logic_section(
        doc,
        "8.3",
        "Single and Multi-Image Sharing",
        "Send cached compressed outputs to other apps through the Android share sheet.",
        [
            "Return immediately when the image list is empty.",
            "For one URI, create ACTION_SEND with the exact image MIME type.",
            "For multiple URIs, create ACTION_SEND_MULTIPLE with image/*.",
            "Attach EXTRA_STREAM values and ClipData entries.",
            "Grant temporary read access with FLAG_GRANT_READ_URI_PERMISSION.",
            "Launch an Android chooser.",
        ],
        "Sharing one WEBP uses ACTION_SEND and image/webp. Sharing three mixed compressed images uses "
        "ACTION_SEND_MULTIPLE and image/* with three stream URIs.",
        ["util/ShareUtils.kt:9-42", "ui/ImageCompressorApp.kt:202-223"],
    )

    add_heading(doc, "9. UI State, Navigation, and Theme Logic", 1)
    add_logic_section(
        doc,
        "9.1",
        "Single Immutable UI State",
        "Keep the Compose interface consistent by rendering one state object rather than many unrelated mutable fields.",
        [
            "Initialize a MutableStateFlow<ImageCompressorUiState> in the ViewModel.",
            "Expose it as read-only StateFlow.",
            "Update state with copy() so each transition is explicit.",
            "Collect state with collectAsStateWithLifecycle() in the root composable.",
            "Render exactly one screen based on state.screen.",
        ],
        "When compression starts, one copy() call sets screen=RESULTS, clears old outputs, sets total, and marks "
        "isCompressing=true. The Results screen immediately renders the progress card.",
        ["ui/ImageCompressorViewModel.kt:37-76", "ui/ImageCompressorViewModel.kt:180-189", "ui/ImageCompressorApp.kt:89-103"],
    )

    add_logic_section(
        doc,
        "9.2",
        "Screen Navigation and Back Mapping",
        "Provide predictable beginner-friendly navigation without a separate navigation framework.",
        [
            "Represent screens with AppScreen enum values.",
            "Update screen directly for top-level navigation.",
            "Map Back from Preview to Home, Settings to Preview, Results to Settings, and History/Settings to Home.",
            "Enable BackHandler for non-root screens.",
            "Show Home, History, and Settings in the bottom navigation bar.",
        ],
        "A user on Results taps Back. The ViewModel maps RESULTS to COMPRESSION_SETTINGS so the user can adjust "
        "quality without reselecting photos.",
        ["ui/ImageCompressorViewModel.kt:22-30", "ui/ImageCompressorViewModel.kt:83-99", "ui/ImageCompressorApp.kt:155-175"],
    )

    add_logic_section(
        doc,
        "9.3",
        "Onboarding and Theme Preferences",
        "Remember first-run completion and the user's chosen appearance across launches.",
        [
            "Read onboarding_complete and theme_preference from private SharedPreferences.",
            "Wrap both values in StateFlow.",
            "Start on ONBOARDING when onboarding_complete is false; otherwise start on HOME.",
            "Update both SharedPreferences and the matching flow when a preference changes.",
            "Resolve SYSTEM, LIGHT, or DARK into a boolean darkTheme value at the Compose root.",
            "Use Android 12+ dynamic color when available, otherwise fall back to app light/dark schemes.",
        ],
        "A first launch sees onboarding. After Get Started, onboarding_complete is saved. If the user later chooses Dark, "
        "the next launch restores DARK and the root theme renders dark colors immediately.",
        ["data/preferences/AppPreferencesRepository.kt:10-51", "ui/ImageCompressorViewModel.kt:56-81", "ui/ImageCompressorApp.kt:89-104", "theme/Theme.kt:27-45"],
    )

    add_logic_section(
        doc,
        "9.4",
        "Responsive Image Grid and UI Statuses",
        "Remain usable across phone widths while clearly showing empty, loading, success, and error states.",
        [
            "Use GridCells.Adaptive(160.dp) so Compose chooses the number of preview columns that fit.",
            "Show LoadingCard while image metadata is loading or compression is running.",
            "Show EmptyState when no selection, no results, or no history exists.",
            "Use SnackbarHostState to present messages, then consume each message after display.",
            "Disable already-saved result buttons and change the label to Saved.",
        ],
        "A narrow phone may fit one 160 dp image card per row; a tablet may fit several. While metadata loads, "
        "the grid is replaced with Reading image details... and a progress indicator.",
        ["ui/ImageCompressorApp.kt:148-153", "ui/ImageCompressorApp.kt:351-418", "ui/ImageCompressorApp.kt:567-654", "ui/ImageCompressorApp.kt:752-789"],
    )

    add_heading(doc, "10. Privacy, Failure Handling, and Constraints", 1)
    add_heading(doc, "10.1 Privacy and Permission Matrix", 2)
    permission_rows = [
        ("Select images", "Android 8.0+", "Photo Picker or document-picker fallback", "No broad gallery-read permission"),
        ("Compress images", "Android 8.0+", "Private cache files", "No runtime permission"),
        ("Save output", "Android 10+", "Scoped MediaStore", "No runtime storage permission"),
        ("Save output", "Android 8/9", "Public Pictures directory", "WRITE_EXTERNAL_STORAGE requested at save time"),
        ("Share output", "Android 8.0+", "FileProvider content URIs", "Temporary URI read grant"),
    ]
    add_table(doc, ["Action", "Android version", "Mechanism", "Permission behavior"], permission_rows, [1400, 1500, 2760, 3700], font_size=8.8)
    add_para(
        doc,
        "The manifest disables backups and excludes root data from cloud backup and device transfer. The privacy onboarding "
        "and Settings screen state that images are compressed locally.",
    )
    add_source(doc, "AndroidManifest.xml:5-41", "res/xml/data_extraction_rules.xml:1-9", "ui/ImageCompressorApp.kt:278-305", "ui/ImageCompressorApp.kt:739-743")

    add_heading(doc, "10.2 Failure Handling", 2)
    failure_rows = [
        ("URI cannot open", "Unable to open the selected image.", "Inspection fails and snackbar receives the error."),
        ("Format bounds unavailable", "This image format could not be decoded.", "Reject unsupported or invalid source."),
        ("Bitmap decode fails", "Unable to decode the selected image.", "Compression stops for the current batch."),
        ("Encode returns false", "Unable to encode the compressed image.", "No cache file is returned."),
        ("MediaStore insert fails", "Unable to create a gallery entry.", "Save action fails clearly."),
        ("Legacy permission denied", "Storage permission was not granted.", "Image remains available in Results but is not saved."),
    ]
    add_table(doc, ["Condition", "Message", "Behavior"], failure_rows, [2200, 3080, 4080], font_size=9)
    add_source(doc, "util/ImageCompressor.kt:134-165", "util/ImageCompressor.kt:258-262", "util/GallerySaver.kt:41-55", "ui/ImageCompressorApp.kt:122-140")

    add_heading(doc, "10.3 Known Constraints and Extension Notes", 2)
    constraints = [
        "Target-size output is best-effort rather than an exact byte guarantee.",
        "Sequential batch processing favors predictable memory use over maximum CPU throughput.",
        "Temporary compressed files live under app cache. Android may evict cache files later; saved gallery copies are durable.",
        "EXIF rotational orientation is handled, but mirrored EXIF orientations are not transformed.",
        "History stores compressed cache URIs and optional gallery URIs. Sharing a very old unsaved cache URI may fail after cache eviction.",
        "Room schema version is 1. Future schema changes should add explicit migrations before release upgrades.",
        "Release minification is disabled. Production publishing can enable R8 after validating sharing, Room, and Compose behavior.",
    ]
    for item in constraints:
        add_bullet(doc, item)

    add_heading(doc, "11. Build, Dependencies, and Verification", 1)
    add_heading(doc, "11.1 Build Configuration", 2)
    build_rows = [
        ("Language", "Kotlin with Java 17 toolchain"),
        ("UI", "Jetpack Compose with Material 3"),
        ("SDK", "compileSdk 37, targetSdk 37, minSdk 26"),
        ("Database", "Room runtime, Room KTX, KSP compiler, exported schema"),
        ("Image preview", "Coil Compose"),
        ("Metadata", "AndroidX ExifInterface"),
        ("Concurrency", "Kotlin Coroutines and Flow through lifecycle/viewmodel dependencies"),
    ]
    add_table(doc, ["Concern", "Configuration"], build_rows, [2200, 7160], font_size=9.5)
    add_source(doc, "app/build.gradle.kts:1-91", "gradle/libs.versions.toml")

    add_heading(doc, "11.2 Automated Checks", 2)
    add_para(doc, "Run the following commands from the project root:")
    add_code(
        doc,
        """
./gradlew testDebugUnitTest assembleDebug
./gradlew lintDebug
""",
    )
    test_rows = [
        ("calculateReductionPercent", "1000 -> 250 bytes", "75%"),
        ("calculateReductionPercent", "0 -> 250 bytes", "0% without division by zero"),
        ("calculateReductionPercent", "100 -> 150 bytes", "0% after negative clamp"),
        ("toReadableSize", "1024 bytes", "1.0 KB"),
        ("toReadableSize", "1,048,576 bytes", "1.00 MB"),
    ]
    add_table(doc, ["Function", "Input", "Expected result"], test_rows, [2500, 3000, 3860], font_size=9.5)
    add_source(doc, "data/model/CompressionModelsTest.kt:1-19", "README.md")

    add_heading(doc, "11.3 Manual Verification Checklist", 2)
    checklist = [
        "Install the debug APK and complete onboarding.",
        "Pick one portrait camera image and confirm its displayed orientation and resolution.",
        "Pick several images and confirm duplicate selections are not repeated.",
        "Compress JPEG at multiple quality values and compare output byte sizes.",
        "Compress with each target-size preset and observe best-effort results.",
        "Test percentage resize and custom resize with aspect-ratio preservation on and off.",
        "Convert a transparent PNG to JPEG and confirm the background becomes white.",
        "Save on Android 10+ and verify the Pictures/Image Compressor album.",
        "Save on Android 8/9 and verify the storage prompt occurs only at save time.",
        "Share one result and several results through the Android share sheet.",
        "Open History, delete one item, clear all items, and switch system/light/dark themes.",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "Appendix A. Pseudocode Reference", 1)
    add_heading(doc, "A.1 Compression Pipeline", 2)
    add_code(
        doc,
        """
compress(image, settings):
    dimensions = calculateOutputDimensions(image, settings)
    bitmap = decodeSampledBitmap(image.uri, dimensions)
    bitmap = scaleBitmap(bitmap, dimensions)
    bitmap = flattenTransparencyForJpeg(bitmap, settings.format)
    bytes = encode(bitmap, settings.format, settings.quality)
    bytes, bitmap = approachTargetSize(bytes, bitmap, settings)
    file = writeUniqueCacheFile(bytes, settings.format)
    normalizeExifOrientation(file)
    recycle(bitmap)
    return CompressedImage(fileProviderUri(file), file.path, file.length, dimensions)
""",
    )
    add_heading(doc, "A.2 Target-Size Loop", 2)
    add_code(
        doc,
        """
attempts = 0
while target exists and encodedBytes > target and attempts < 40:
    if format is lossy and quality > 10:
        quality = max(10, quality - 5)
    else:
        bitmap = scale(bitmap, width * 0.9, height * 0.9)
    encodedBytes = encode(bitmap, format, quality)
    attempts += 1
""",
    )
    add_heading(doc, "A.3 Pixel-Budget Constraint", 2)
    add_code(
        doc,
        """
pixels = width * height
if pixels <= 16_000_000:
    return width, height
scale = sqrt(16_000_000 / pixels)
return round(width * scale), round(height * scale)
""",
    )
    add_heading(doc, "A.4 MediaStore Save", 2)
    add_code(
        doc,
        """
if Android >= 10:
    uri = insert MediaStore row with IS_PENDING = 1
    copy cache file into uri output stream
    update row with IS_PENDING = 0
else:
    require WRITE_EXTERNAL_STORAGE
    copy into Pictures/Image Compressor
    scan file with MediaScannerConnection
""",
    )

    add_heading(doc, "Appendix B. Source-File Ownership Map", 1)
    source_rows = [
        ("MainActivity.kt", "Compose entry point and edge-to-edge activity setup"),
        ("ImageCompressorApplication.kt", "Application-scoped container and dependency construction"),
        ("ui/ImageCompressorApp.kt", "All Compose screens, Photo Picker, save permission launcher, status UI"),
        ("ui/ImageCompressorViewModel.kt", "MVVM state, actions, progress, save/history orchestration"),
        ("data/model/CompressionModels.kt", "Domain models, enums, reduction and readable-size helpers"),
        ("data/ImageRepository.kt", "Repository interface and implementation for batch coordination"),
        ("data/local/*", "Room entity, DAO, and database"),
        ("data/preferences/*", "Onboarding and theme SharedPreferences flows"),
        ("util/ImageCompressor.kt", "Bitmap inspection, resizing, sampling, EXIF handling, encoding, cache output"),
        ("util/GallerySaver.kt", "Android-version-specific gallery publishing"),
        ("util/ShareUtils.kt", "Single and multi-image Android sharing"),
        ("AndroidManifest.xml", "Legacy save permission, FileProvider, Photo Picker backport hook"),
        ("res/xml/file_paths.xml", "FileProvider cache-path exposure"),
        ("res/xml/data_extraction_rules.xml", "Backup and device-transfer exclusions"),
        ("app/build.gradle.kts", "SDK levels, plugins, and dependencies"),
        ("data/model/CompressionModelsTest.kt", "Unit tests for derived size math"),
    ]
    add_table(doc, ["Source file", "Ownership"], source_rows, [3400, 5960], font_size=9)

    doc.add_page_break()
    add_heading(doc, "Appendix C. Practical Extension Playbook", 1)
    extension_rows = [
        ("Add HEIC output", "Extend OutputFormat, confirm Bitmap encoder support, update MIME mapping and UI chips."),
        ("Parallelize batches", "Use bounded concurrency and a semaphore; measure heap behavior before increasing parallelism."),
        ("Exact target sizing", "Replace linear quality steps with a bounded binary search, then dimension fallback."),
        ("Retain durable output", "Move unsaved history outputs from cache to app files storage or add explicit export policy."),
        ("Add Room schema fields", "Increment database version and ship a migration instead of relying on destructive fallback."),
        ("Add mirrored EXIF support", "Map flip and transpose EXIF values into Matrix transforms before export."),
    ]
    add_table(doc, ["Change", "Implementation direction"], extension_rows, [2600, 6760], font_size=9.5)

    add_callout(
        doc,
        "Closing note",
        "This handbook describes the implemented project as of June 2, 2026. Keep it versioned with the source code and "
        "update the algorithm examples whenever compression constants, permissions, or persistence behavior change.",
        fill="EAF5EE",
        label_color="1F5B36",
    )

    core = doc.core_properties
    core.title = "Image Compressor Technical Documentation"
    core.subject = "Architecture, logic, algorithms, and worked examples"
    core.author = "Codex"
    core.keywords = "Android, Kotlin, Jetpack Compose, image compression, Room, MediaStore"
    core.comments = "Generated from the implemented Image Compressor project."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
